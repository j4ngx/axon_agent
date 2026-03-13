"""Telegram message handlers and authorization middleware.

All Telegram-specific logic (routing, filtering, auth) lives here and
delegates business work to the agent loop.
"""

from __future__ import annotations

import io
import logging
from collections.abc import Awaitable, Callable
from typing import Any

import fitz  # PyMuPDF
from aiogram import BaseMiddleware, Bot, Router, types
from aiogram.filters import Command, CommandStart
from aiogram.types import BufferedInputFile
from docx import Document as DocxDocument

from helix.agent.loop import AgentLoop
from helix.llm.embeddings import EmbeddingClient
from helix.llm.transcription import TranscriptionClient
from helix.llm.tts import TTSClient
from helix.llm.vision import VisionClient
from helix.memory.document_repository import DocumentRepository
from helix.memory.models import Document, DocumentChunk, VoiceNote
from helix.memory.voice_note_repository import VoiceNoteRepository

logger = logging.getLogger(__name__)

_SUPPORTED_MIMES = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "text/plain",
}


# ---------------------------------------------------------------------------
# Authorization middleware
# ---------------------------------------------------------------------------


class AuthMiddleware(BaseMiddleware):
    """Reject messages from users not in the allow-list.

    Args:
        allowed_user_ids: Set of Telegram user IDs that are authorised.
    """

    def __init__(self, allowed_user_ids: set[int]) -> None:
        super().__init__()
        self._allowed = allowed_user_ids

    async def __call__(
        self,
        handler: Callable[[types.Update, dict[str, Any]], Awaitable[Any]],
        event: types.Message,
        data: dict[str, Any],
    ) -> Any:
        """Check authorisation before forwarding to the actual handler."""
        user = event.from_user
        if user is None or user.id not in self._allowed:
            uid = user.id if user else "unknown"
            logger.warning("Unauthorized access attempt", extra={"user_id": uid})
            await event.answer("⛔ You are not authorised to use this bot.")
            return None
        return await handler(event, data)


# ---------------------------------------------------------------------------
# Router factory
# ---------------------------------------------------------------------------


def create_router(
    agent_loop: AgentLoop,
    allowed_user_ids: list[int],
    transcription_client: TranscriptionClient,
    bot: Bot,
    vision_client: VisionClient | None = None,
    voice_note_repo: VoiceNoteRepository | None = None,
    document_repo: DocumentRepository | None = None,
    embedding_client: EmbeddingClient | None = None,
    tts_client: TTSClient | None = None,
) -> Router:
    """Build the aiogram ``Router`` with auth middleware and handlers.

    Args:
        agent_loop: The agent loop that processes user messages.
        allowed_user_ids: Telegram user IDs allowed to interact with Helix.
        transcription_client: Groq Whisper client for voice messages.
        bot: The aiogram Bot instance (needed to fetch file URLs).
        vision_client: Optional Groq Vision client for photo descriptions.
        voice_note_repo: Optional repo to auto-save voice transcriptions.
        document_repo: Optional repo for document storage.
        embedding_client: Optional embedding client for document chunks.
        tts_client: Optional ElevenLabs TTS client for audio responses.

    Returns:
        A configured ``Router``.
    """
    router = Router(name="helix_main")
    router.message.middleware(AuthMiddleware(set(allowed_user_ids)))

    # Per-user in-memory toggle for voice mode (text → audio).
    _voice_mode_users: set[int] = set()

    # -- /start command -------------------------------------------------

    @router.message(CommandStart())
    async def handle_start(message: types.Message) -> None:
        """Greet the user on ``/start``."""
        await message.answer(
            "👋 Hi! I'm **Helix**, your personal AI assistant.\n\n"
            "Send me any message — text or a voice note — and I'll do my best to help.",
            parse_mode="Markdown",
        )

    # -- /voice toggle command ------------------------------------------

    @router.message(Command("voice"))
    async def handle_voice_toggle(message: types.Message) -> None:
        """Toggle voice mode — audio responses for text messages."""
        user_id = message.from_user.id  # type: ignore[union-attr]

        if tts_client is None:
            await message.answer("⚠️ Voice responses are not configured.")
            return

        if user_id in _voice_mode_users:
            _voice_mode_users.discard(user_id)
            await message.answer("🔇 Voice mode disabled.")
        else:
            _voice_mode_users.add(user_id)
            await message.answer("🔊 Voice mode enabled — I'll reply with audio too.")

    # -- Voice / audio messages -----------------------------------------

    @router.message(lambda m: m.voice is not None or m.audio is not None)
    async def handle_voice(message: types.Message) -> None:
        """Transcribe a voice note or audio file and run it through the agent."""
        user_id = message.from_user.id  # type: ignore[union-attr]

        # Determine the Telegram file object (voice note vs. audio file).
        file_obj = message.voice or message.audio
        if file_obj is None:
            await message.answer("⚠️ Could not read the audio file.")
            return

        logger.info(
            "Incoming voice message",
            extra={"user_id": user_id, "file_id": file_obj.file_id},
        )

        # Show the user we're working on it.
        await message.answer("🎙️ Transcribing your voice message…")

        try:
            # 1. Get the download URL from Telegram.
            file_info = await bot.get_file(file_obj.file_id)
            if file_info.file_path is None:
                await message.answer("⚠️ Could not retrieve the audio file from Telegram.")
                return

            # Determine filename/extension for codec detection.
            file_name = getattr(message.audio, "file_name", None) or f"voice_{file_obj.file_id}.ogg"
            file_url = bot.session.api.file_url(
                bot.token,
                file_info.file_path,
            )

            # 2. Transcribe.
            transcribed_text = await transcription_client.transcribe_from_url(
                file_url=file_url,
                file_name=file_name,
            )
        except Exception:
            logger.exception("Voice transcription failed", extra={"user_id": user_id})
            await message.answer(
                "⚠️ Sorry, I couldn't transcribe that audio. "
                "Please try again or send a text message."
            )
            return

        logger.info(
            "Voice transcribed",
            extra={"user_id": user_id, "text_length": len(transcribed_text)},
        )

        # Echo the transcription so the user can confirm.
        # Escape Markdown special chars to avoid Telegram parse errors.
        safe_text = _escape_markdown(transcribed_text)
        await message.answer(f"🗣️ *Transcription:*\n_{safe_text}_", parse_mode="Markdown")

        # 2b. Auto-save voice note to Firestore.
        if voice_note_repo:
            try:
                duration = getattr(file_obj, "duration", 0) or 0
                note = VoiceNote(
                    user_id=user_id,
                    text=transcribed_text,
                    duration_seconds=duration,
                    telegram_file_id=file_obj.file_id,
                )
                await voice_note_repo.create(note)
            except Exception:
                logger.exception("Failed to save voice note", extra={"user_id": user_id})

        # 3. Run through the agent loop as if it were a text message.
        try:
            reply = await agent_loop.run(user_id=user_id, user_message=transcribed_text)
        except Exception:
            logger.exception("Agent loop error after transcription", extra={"user_id": user_id})
            reply = "Something went wrong while processing your message. Please try again."

        for chunk in _split_message(reply):
            await _safe_reply(message, chunk)

        # Voice messages always get an audio reply when TTS is available.
        await _send_audio_reply(message, tts_client, reply)

    # -- Photo messages -------------------------------------------------

    @router.message(lambda m: m.photo is not None and m.photo)
    async def handle_photo(message: types.Message) -> None:
        """Describe a photo via Groq Vision and pass to the agent."""
        user_id = message.from_user.id  # type: ignore[union-attr]

        if vision_client is None:
            await message.answer("⚠️ Image processing is not configured.")
            return

        logger.info("Incoming photo", extra={"user_id": user_id})
        await message.answer("🖼️ Analysing your image…")

        try:
            # Download the largest available photo size.
            photo = message.photo[-1]
            file_info = await bot.get_file(photo.file_id)
            if file_info.file_path is None:
                await message.answer("⚠️ Could not retrieve the image from Telegram.")
                return

            file_url = bot.session.api.file_url(bot.token, file_info.file_path)
            import httpx

            async with httpx.AsyncClient(timeout=30.0) as http:
                resp = await http.get(file_url)
                resp.raise_for_status()
                image_bytes = resp.content

            # Detect MIME type from file path extension.
            ext = (file_info.file_path or "").rsplit(".", 1)[-1].lower()
            mime_map = {"jpg": "image/jpeg", "jpeg": "image/jpeg", "png": "image/png"}
            mime_type = mime_map.get(ext, "image/jpeg")

            description = await vision_client.describe_image(image_bytes, mime_type)
        except Exception:
            logger.exception("Photo processing failed", extra={"user_id": user_id})
            await message.answer("⚠️ Sorry, I couldn't process that image.")
            return

        # Pass the description through the agent loop with optional caption.
        caption = message.caption or ""
        agent_input = f"[Photo description: {description}]"
        if caption:
            agent_input += f"\nUser caption: {caption}"

        try:
            reply = await agent_loop.run(user_id=user_id, user_message=agent_input)
        except Exception:
            logger.exception("Agent loop error after photo", extra={"user_id": user_id})
            reply = "Something went wrong while processing your image."

        for chunk in _split_message(reply):
            await _safe_reply(message, chunk)

    # -- Document uploads -----------------------------------------------

    @router.message(lambda m: m.document is not None)
    async def handle_document(message: types.Message) -> None:
        """Extract text from uploaded documents, chunk, embed, and store."""
        user_id = message.from_user.id  # type: ignore[union-attr]

        if document_repo is None or embedding_client is None:
            await message.answer("⚠️ Document processing is not configured.")
            return

        doc_obj = message.document
        if doc_obj is None:
            await message.answer("⚠️ Could not read the document.")
            return

        mime = doc_obj.mime_type or ""
        if mime not in _SUPPORTED_MIMES:
            await message.answer("⚠️ Unsupported file type. I can process PDF, DOCX, and TXT files.")
            return

        logger.info(
            "Incoming document",
            extra={"user_id": user_id, "mime": mime, "filename": doc_obj.file_name},
        )
        await message.answer("📄 Processing your document…")

        try:
            # Download the file.
            file_info = await bot.get_file(doc_obj.file_id)
            if file_info.file_path is None:
                await message.answer("⚠️ Could not retrieve the file from Telegram.")
                return

            file_url = bot.session.api.file_url(bot.token, file_info.file_path)
            import httpx

            async with httpx.AsyncClient(timeout=60.0) as http:
                resp = await http.get(file_url)
                resp.raise_for_status()
                file_bytes = resp.content

            # Extract text.
            text, page_count = _extract_text(file_bytes, mime)
            if not text.strip():
                await message.answer("⚠️ The document appears to be empty or unreadable.")
                return

            # Chunk the text.
            chunks_text = _chunk_text(text)

            # Store document metadata.
            doc_model = Document(
                user_id=user_id,
                filename=doc_obj.file_name or "unknown",
                mime_type=mime,
                page_count=page_count,
                chunk_count=len(chunks_text),
            )
            doc_id = await document_repo.create_document(doc_model)

            # Embed and store chunks.
            embeddings = await embedding_client.embed(chunks_text)
            chunk_models = [
                DocumentChunk(
                    user_id=user_id,
                    document_id=doc_id,
                    text=ct,
                    chunk_index=i,
                    embedding=embeddings[i] if i < len(embeddings) else None,
                )
                for i, ct in enumerate(chunks_text)
            ]
            await document_repo.save_chunks(chunk_models)

            await message.answer(
                f"✅ **{doc_obj.file_name}** processed — "
                f"{len(chunks_text)} chunks, {page_count} pages.\n"
                "You can now ask questions about it using the `document_qa` tool.",
                parse_mode="Markdown",
            )
        except Exception:
            logger.exception("Document processing failed", extra={"user_id": user_id})
            await message.answer("⚠️ Sorry, I couldn't process that document.")

    # -- Text messages --------------------------------------------------

    @router.message()
    async def handle_message(message: types.Message) -> None:
        """Route text messages through the agent loop."""
        if not message.text:
            await message.answer(
                "I can only process text and voice messages. "
                "Please send a text or record a voice note 🎙️"
            )
            return

        user_id = message.from_user.id  # type: ignore[union-attr]
        logger.info(
            "Incoming message",
            extra={"user_id": user_id, "length": len(message.text)},
        )

        try:
            reply = await agent_loop.run(user_id=user_id, user_message=message.text)
        except Exception:
            logger.exception("Agent loop error", extra={"user_id": user_id})
            reply = "Something went wrong while processing your message. Please try again."

        # Telegram has a 4096-char limit per message.
        for chunk in _split_message(reply):
            await _safe_reply(message, chunk)

        # Send audio reply if user has voice mode enabled.
        if user_id in _voice_mode_users:
            await _send_audio_reply(message, tts_client, reply)

    return router


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

_MARKDOWN_SPECIAL = r"\_*[]()~`>#+-=|{}.!"


def _escape_markdown(text: str) -> str:
    """Escape Telegram Markdown v1 special characters.

    Args:
        text: Raw text that may contain special chars.

    Returns:
        Escaped text safe for ``parse_mode="Markdown"``.
    """
    for ch in ("_", "*", "`", "["):
        text = text.replace(ch, f"\\{ch}")
    return text


async def _safe_reply(message: types.Message, text: str) -> None:
    """Send *text* with Markdown, falling back to plain text on parse errors."""
    try:
        await message.answer(text, parse_mode="Markdown")
    except Exception:
        await message.answer(text)


_MAX_MESSAGE_LENGTH = 4096


def _split_message(text: str, max_length: int = _MAX_MESSAGE_LENGTH) -> list[str]:
    """Split a long message into chunks that fit Telegram's limit.

    Args:
        text: The full reply text.
        max_length: Maximum characters per chunk.

    Returns:
        A list of one or more message chunks.
    """
    if len(text) <= max_length:
        return [text]
    chunks: list[str] = []
    while text:
        chunks.append(text[:max_length])
        text = text[max_length:]
    return chunks


async def _send_audio_reply(
    message: types.Message,
    tts_client: TTSClient | None,
    text: str,
) -> None:
    """Synthesise *text* and send it as an MP3 audio message.

    Silently skips if *tts_client* is ``None`` or synthesis fails —
    the text reply has already been sent.
    """
    if tts_client is None:
        return
    try:
        audio_bytes = await tts_client.synthesize(text)
        audio_file = BufferedInputFile(audio_bytes, filename="helix.mp3")
        await message.answer_audio(audio=audio_file)
    except Exception:
        logger.warning("TTS audio reply failed — text was already sent", exc_info=True)


# ---------------------------------------------------------------------------
# Document extraction helpers
# ---------------------------------------------------------------------------

_CHUNK_SIZE = 800
_CHUNK_OVERLAP = 100


def _extract_text(file_bytes: bytes, mime_type: str) -> tuple[str, int]:
    """Extract text from a file based on its MIME type.

    Returns:
        A tuple of ``(text, page_count)``.
    """
    if mime_type == "application/pdf":
        return _extract_text_from_pdf(file_bytes)
    if mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return _extract_text_from_docx(file_bytes)
    # text/plain
    return file_bytes.decode("utf-8", errors="replace"), 0


def _extract_text_from_pdf(file_bytes: bytes) -> tuple[str, int]:
    """Extract text from a PDF using PyMuPDF."""
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    pages: list[str] = []
    for page in doc:
        pages.append(page.get_text())
    doc.close()
    return "\n\n".join(pages), len(pages)


def _extract_text_from_docx(file_bytes: bytes) -> tuple[str, int]:
    """Extract text from a DOCX using python-docx."""
    doc = DocxDocument(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    return "\n\n".join(paragraphs), 0


def _chunk_text(
    text: str,
    chunk_size: int = _CHUNK_SIZE,
    overlap: int = _CHUNK_OVERLAP,
) -> list[str]:
    """Split text into overlapping chunks for embedding.

    Args:
        text: The full document text.
        chunk_size: Target characters per chunk.
        overlap: Number of overlapping characters between chunks.

    Returns:
        A list of text chunks.
    """
    if len(text) <= chunk_size:
        return [text]

    chunks: list[str] = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunks.append(text[start:end])
        start = end - overlap

    return chunks
